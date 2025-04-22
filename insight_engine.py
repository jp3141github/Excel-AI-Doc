from openai import OpenAI
import pandas as pd
from typing import Dict, Any
import io
from datetime import datetime
from docx import Document
import os
import openpyxl  # Required for Excel reading

# ✅ Instantiate OpenAI client using your API key
client = OpenAI(api_key="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")

def load_excel_sheet(file_content: bytes, sheet_name: str = None) -> pd.DataFrame:
    """Loads an Excel sheet from a byte stream."""
    try:
        return pd.read_excel(io.BytesIO(file_content), sheet_name=sheet_name)
    except Exception as e:
        raise ValueError(f"Failed to load Excel sheet: {e}")

def generate_summary_stats(df: pd.DataFrame) -> Dict[str, Any]:
    """Generates descriptive statistics and metadata for each column."""
    try:
        desc = df.describe(include='all').to_dict()
        info = {
            col: {
                "dtype": str(df[col].dtype),
                "nulls": int(df[col].isnull().sum()),
                "unique": int(df[col].nunique()),
                "sample_values": df[col].dropna().unique()[:3].tolist()
            }
            for col in df.columns
        }
        return {"stats": desc, "structure": info}
    except Exception as e:
        raise RuntimeError(f"Error generating summary stats: {e}")

def detect_quality_issues(df: pd.DataFrame) -> Dict[str, str]:
    """Detects columns with potential data quality issues."""
    issues = {}
    for col in df.columns:
        null_ratio = df[col].isnull().mean()
        unique_count = df[col].nunique(dropna=True)
        if null_ratio > 0.5:
            issues[col] = "More than 50% missing values"
        elif unique_count == 1:
            issues[col] = "Single unique value (likely static column)"
    return issues

def build_gpt_prompt(stats: dict, structure: dict, issues: dict) -> str:
    """Constructs a prompt for GPT-4 to generate insights."""
    return f"""
You are a highly skilled data analyst, data scientist and technical writer. Your task is to examine the dataset and generate a professional, structured insight report.

Your responsibilities include:

1. Generate meaningful high-level overview/summaries of the dataset
2. Providing column-by-column commentary on structure and contents
3. Identifying trends, anomalies, outliers, and significant patterns
4. Highlighting relationships and correlations across key variables
5. Recommending visualizations (e.g., histograms, time series, heatmaps) and statistical techniques (e.g., regression, clustering, correlation analysis) for deeper analysis
6. Detecting data quality concerns and limitations
7. Suggesting relevant data aggregations or transformations
8. Provide recommendations for how this dataset can be used in reporting or analytics
9. Creating and displaying examples of any recommendations, especially visualisations (charts and summary tables)

In addition, consider the following use cases for this dataset:
- Predicting outcomes based on historical features
- Analyzing performance by different cohorts over time
- Identifying emerging trends and behavioral patterns by cohort
- Supporting actuarial, pricing, and portfolio strategy initiatives

You are provided with:
- Summary statistics for numerical and categorical columns
- Metadata on column types, nulls, unique values, and sample entries
- Flagged data quality issues such as high nulls or suspiciously constant fields

Use this to produce a structured written report output, including:
- Dataset overview
- Column-by-column commentary / descriptions
- Key observations, patterns and correlations
- Data quality concerns
- Suggestions for improvement / Recommendations for Next Steps

Summary Stats:
{stats}

Column Metadata:
{structure}

Detected Issues:
{issues}
"""

def query_openai_insights(prompt: str) -> str:
    """Queries OpenAI's GPT-4 API with real-time streaming output."""
    try:
        response_stream = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            stream=True  # ✅ Enables real-time stream
        )

        collected_chunks = []
        print("\n🧠 GPT-4 Insight Output Preview (streaming):\n")

        for chunk in response_stream:
            delta = chunk.choices[0].delta
            if delta.content:
                print(delta.content, end="", flush=True)  # ✅ Realtime print
                collected_chunks.append(delta.content)

        full_text = "".join(collected_chunks)
        print("\n\n✅ Streaming complete.\n")
        return full_text

    except Exception as e:
        raise RuntimeError(f"Failed to query GPT-4: {e}")

def export_report_to_word(prompt: str, stats: dict, structure: dict, issues: dict, insights: str) -> str:
    """Exports the full report to a Word document with a timestamped filename."""
    doc = Document()
    doc.add_heading("Data Insight Report", 0)

    doc.add_heading("Prompt Sent to GPT", level=1)
    doc.add_paragraph(prompt)

    doc.add_heading("Summary Stats", level=1)
    doc.add_paragraph(str(stats))

    doc.add_heading("Column Metadata", level=1)
    doc.add_paragraph(str(structure))

    doc.add_heading("Detected Issues", level=1)
    doc.add_paragraph(str(issues))

    doc.add_heading("GPT-4 Insights", level=1)
    doc.add_paragraph(insights)

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    file_path = f"insight_report_{timestamp}.docx"
    doc.save(file_path)
    return file_path
